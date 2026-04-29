"""
CampusRule AI v2.0 — FastAPI Server.

Full-featured API with chat (streaming + non-streaming), autocomplete,
conversation history, analytics, feedback, and document management.
"""

import json
import time
import uuid
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Optional

from fastapi import BackgroundTasks, FastAPI, File, HTTPException, UploadFile, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel, Field
import shutil

from models import (
    ChatRequest, ChatResponse,
    AutocompleteRequest, AutocompleteResponse,
    SearchRequest, SearchResponse,
    FeedbackRequest, FeedbackResponse,
    ExportRequest,
    AnalyticsResponse,
)
from retrieval import (
    DOCS_DIR,
    build_index,
    get_all_documents,
    get_chunks_for_context,
    hybrid_search,
    load_index,
    search,
)
from rag import (
    detect_intent,
    extract_entities,
    generate_answer,
    generate_answer_stream,
    check_ollama_status,
)
from predictor import get_engine
from conversation import (
    add_message,
    create_conversation,
    delete_conversation,
    get_conversation,
    get_conversation_history,
    list_conversations,
    pin_conversation,
)
from analytics import (
    get_analytics,
    save_feedback,
    track_query,
)


@asynccontextmanager
async def lifespan(app: FastAPI):
    if not load_index():
        print("No existing index found — building from PDFs...")
        build_index()
    # Initialize autocomplete engine
    get_engine()
    yield


app = FastAPI(
    title="CampusRule AI",
    description="Advanced AI chatbot for university academic policies with RAG, "
                "semantic search, and local LLM integration.",
    version="2.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000", "http://127.0.0.1:5173", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ══════════════════════════════════════════════════════════════════════════════
#  HEALTH & STATUS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/", tags=["Health"])
async def root():
    return {"message": "CampusRule AI v2.0 is running", "status": "ok", "version": "2.0.0"}


@app.get("/api/health", tags=["Health"])
async def health():
    ollama = await check_ollama_status()
    return {
        "status": "healthy",
        "version": "2.0.0",
        "ollama": ollama,
    }


# ══════════════════════════════════════════════════════════════════════════════
#  CHAT (Core Feature)
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/chat", tags=["Chat"])
async def chat_endpoint(request: ChatRequest, background_tasks: BackgroundTasks):
    """
    Main chat endpoint. Supports streaming (SSE) and non-streaming responses.
    """
    query = request.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty.")

    conversation_id = request.conversation_id or str(uuid.uuid4())
    create_conversation(conversation_id, request.user_id)

    # Save user message
    add_message(conversation_id, "user", query)

    # Detect intent
    intent = detect_intent(query)

    # Retrieve relevant chunks
    chunks, context = get_chunks_for_context(query, top_k=5)

    if not context:
        response = ChatResponse(
            answer="I couldn't find any relevant policy information for your question. "
                   "Please try rephrasing or ask about a specific academic policy.",
            confidence=0.0,
            conversation_id=conversation_id,
            intent=intent,
        )
        add_message(conversation_id, "assistant", response.answer, confidence=0.0, intent=intent)
        return response

    # Get conversation history for context
    history = get_conversation_history(conversation_id, limit=6)

    if request.stream:
        # Streaming response (Server-Sent Events)
        async def event_stream():
            full_answer = []
            async for chunk in generate_answer_stream(query, context, chunks, history, intent):
                yield chunk
                # Collect tokens for saving
                try:
                    line = chunk.strip()
                    if line.startswith("data: "):
                        data = json.loads(line[6:])
                        if data.get("type") == "token":
                            full_answer.append(data["content"])
                except Exception:
                    pass

            # Save assistant message after streaming completes
            answer_text = "".join(full_answer)
            if answer_text:
                confidence = 0.0
                if chunks:
                    from rag import calculate_confidence
                    confidence = calculate_confidence(chunks)

                add_message(
                    conversation_id, "assistant", answer_text,
                    sources=[{"document": c["document"], "section": c["section"], "page": c["page"]} for c in chunks[:3]],
                    confidence=confidence, intent=intent,
                )

                # Track analytics
                background_tasks.add_task(
                    track_query, query,
                    category=chunks[0]["category"] if chunks else "general",
                    confidence=confidence,
                    intent=intent,
                )

        return StreamingResponse(
            event_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )
    else:
        # Non-streaming response
        start_time = time.time()
        result = await generate_answer(query, context, chunks, history, intent)
        thinking_time = round(time.time() - start_time, 2)

        # Save assistant message
        add_message(
            conversation_id, "assistant", result["answer"],
            sources=result.get("sources", []),
            confidence=result.get("confidence", 0.0),
            follow_ups=result.get("follow_ups", []),
            intent=intent,
        )

        # Track analytics
        background_tasks.add_task(
            track_query, query,
            category=chunks[0]["category"] if chunks else "general",
            confidence=result.get("confidence", 0.0),
            response_time=thinking_time,
            intent=intent,
        )

        response = ChatResponse(
            answer=result["answer"],
            sources=[
                {
                    "text": c["text"][:200],
                    "document": c["document"],
                    "page": c["page"],
                    "category": c["category"],
                    "section": c["section"],
                    "score": c["score"],
                }
                for c in chunks
            ],
            confidence=result.get("confidence", 0.0),
            follow_ups=result.get("follow_ups", []),
            thinking_time=thinking_time,
            intent=intent,
            conversation_id=conversation_id,
            cited_documents=result.get("cited_documents", []),
        )
        return response


# ══════════════════════════════════════════════════════════════════════════════
#  AUTOCOMPLETE
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/autocomplete", tags=["Autocomplete"])
async def autocomplete_endpoint(request: AutocompleteRequest):
    """Get autocomplete suggestions for a query prefix."""
    start = time.time()
    engine = get_engine()
    suggestions = engine.suggest(request.prefix, limit=request.limit, context=request.context)
    elapsed_ms = round((time.time() - start) * 1000, 1)

    return AutocompleteResponse(
        suggestions=suggestions,
        query_time_ms=elapsed_ms,
    )


@app.post("/api/autocomplete/click", tags=["Autocomplete"])
async def autocomplete_click(phrase: str = Query(...)):
    """Record a click on a suggestion (learning signal)."""
    engine = get_engine()
    engine.record_click(phrase)
    return {"status": "ok"}


# ══════════════════════════════════════════════════════════════════════════════
#  SEARCH (Legacy / Backward Compatible)
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/search", response_model=SearchResponse, tags=["Search"])
async def search_endpoint(request: SearchRequest):
    query = request.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Query cannot be empty.")

    result = search(query, top_k=request.top_k)
    if "error" in result:
        raise HTTPException(status_code=500, detail=result["error"])
    return result


# ══════════════════════════════════════════════════════════════════════════════
#  CONVERSATIONS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/conversations", tags=["Conversations"])
async def list_conversations_endpoint(
    user_id: Optional[str] = None,
    limit: int = Query(default=20, ge=1, le=100),
):
    """List recent conversations."""
    return {"conversations": list_conversations(user_id, limit)}


@app.get("/api/conversation/{conversation_id}", tags=["Conversations"])
async def get_conversation_endpoint(conversation_id: str):
    """Get full conversation history."""
    conv = get_conversation(conversation_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found.")
    return conv


@app.post("/api/conversation/{conversation_id}/pin", tags=["Conversations"])
async def pin_conversation_endpoint(conversation_id: str, pinned: bool = True):
    """Pin or unpin a conversation."""
    pin_conversation(conversation_id, pinned)
    return {"status": "ok", "pinned": pinned}


@app.delete("/api/conversation/{conversation_id}", tags=["Conversations"])
async def delete_conversation_endpoint(conversation_id: str):
    """Delete a conversation."""
    delete_conversation(conversation_id)
    return {"status": "ok", "message": "Conversation deleted."}


# ══════════════════════════════════════════════════════════════════════════════
#  FEEDBACK
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/feedback", response_model=FeedbackResponse, tags=["Feedback"])
async def feedback_endpoint(request: FeedbackRequest):
    """Submit feedback on a response."""
    save_feedback(
        message_id=request.message_id,
        conversation_id=request.conversation_id,
        rating=request.rating,
        helpful=request.helpful,
        comment=request.comment or "",
    )
    return FeedbackResponse()


# ══════════════════════════════════════════════════════════════════════════════
#  ANALYTICS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/analytics", tags=["Analytics"])
async def analytics_endpoint(days: int = Query(default=30, ge=1, le=365)):
    """Get analytics dashboard data."""
    return get_analytics(days)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/documents", tags=["Documents"])
async def list_documents():
    docs = get_all_documents()
    categories = {}
    for doc in docs:
        cat = doc["category"]
        categories[cat] = categories.get(cat, 0) + 1
    return {
        "documents": docs,
        "total": len(docs),
        "categories": categories,
    }


@app.post("/api/upload", tags=["Documents"])
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
):
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted.")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    dest = DOCS_DIR / file.filename

    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    background_tasks.add_task(build_index)

    return {
        "message": f"Uploaded '{file.filename}'. Index is rebuilding in the background.",
        "filename": file.filename,
    }


@app.delete("/api/documents/{filename}", tags=["Documents"])
async def delete_document(filename: str, background_tasks: BackgroundTasks):
    target = DOCS_DIR / filename
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"Document '{filename}' not found.")

    target.unlink()
    background_tasks.add_task(build_index)
    return {"message": f"Deleted '{filename}'. Index is rebuilding in the background."}


@app.post("/api/admin/reindex", tags=["Admin"])
async def reindex(background_tasks: BackgroundTasks):
    """Force rebuild the FAISS index from all PDFs."""
    background_tasks.add_task(build_index)
    return {"message": "Index rebuild started in the background."}


# ══════════════════════════════════════════════════════════════════════════════
#  EXPORT
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/export", tags=["Export"])
async def export_conversation(request: ExportRequest):
    """Export a conversation in the requested format."""
    conv = get_conversation(request.conversation_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found.")

    export_dir = Path(__file__).parent / "data" / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)

    filename = f"campusrule_{request.conversation_id[:8]}"

    if request.format == "txt":
        content = _export_text(conv)
        filepath = export_dir / f"{filename}.txt"
        filepath.write_text(content)
        return FileResponse(filepath, filename=f"{filename}.txt", media_type="text/plain")

    elif request.format == "md":
        content = _export_markdown(conv)
        filepath = export_dir / f"{filename}.md"
        filepath.write_text(content)
        return FileResponse(filepath, filename=f"{filename}.md", media_type="text/markdown")

    elif request.format == "docx":
        filepath = _export_docx(conv, export_dir, filename)
        return FileResponse(filepath, filename=f"{filename}.docx",
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif request.format == "pdf":
        filepath = _export_pdf(conv, export_dir, filename)
        return FileResponse(filepath, filename=f"{filename}.pdf", media_type="application/pdf")

    raise HTTPException(status_code=400, detail=f"Unsupported format: {request.format}")


def _export_text(conv: dict) -> str:
    lines = [
        f"CampusRule AI — Conversation Export",
        f"Title: {conv.get('title', 'Untitled')}",
        f"Date: {conv.get('created_at', '')}",
        f"{'=' * 60}",
        "",
    ]
    for msg in conv.get("messages", []):
        role = "🧑 Student" if msg["role"] == "user" else "🤖 CampusRule AI"
        lines.append(f"{role} ({msg['timestamp']}):")
        lines.append(msg["content"])
        if msg.get("sources"):
            lines.append(f"  Sources: {', '.join(s.get('document', '') for s in msg['sources'])}")
        lines.append("")
    return "\n".join(lines)


def _export_markdown(conv: dict) -> str:
    lines = [
        f"# CampusRule AI — Conversation Export",
        f"",
        f"**Title:** {conv.get('title', 'Untitled')}  ",
        f"**Date:** {conv.get('created_at', '')}",
        f"",
        f"---",
        f"",
    ]
    for msg in conv.get("messages", []):
        if msg["role"] == "user":
            lines.append(f"### 🧑 Student")
        else:
            lines.append(f"### 🤖 CampusRule AI")
        lines.append(f"*{msg['timestamp']}*")
        lines.append(f"")
        lines.append(msg["content"])
        if msg.get("sources"):
            lines.append(f"")
            lines.append(f"> **Sources:** {', '.join(s.get('document', '') for s in msg['sources'])}")
        lines.append(f"")
        lines.append(f"---")
        lines.append(f"")
    return "\n".join(lines)


def _export_docx(conv: dict, export_dir: Path, filename: str) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()
        doc.add_heading("CampusRule AI — Conversation Export", 0)
        doc.add_paragraph(f"Title: {conv.get('title', 'Untitled')}")
        doc.add_paragraph(f"Date: {conv.get('created_at', '')}")

        for msg in conv.get("messages", []):
            role = "Student" if msg["role"] == "user" else "CampusRule AI"
            p = doc.add_heading(f"{role}", level=2)
            doc.add_paragraph(msg["content"])
            if msg.get("sources"):
                sources_text = ", ".join(s.get("document", "") for s in msg["sources"])
                doc.add_paragraph(f"Sources: {sources_text}").italic = True

        filepath = export_dir / f"{filename}.docx"
        doc.save(str(filepath))
        return filepath
    except ImportError:
        # Fallback to text
        filepath = export_dir / f"{filename}.txt"
        filepath.write_text(_export_text(conv))
        return filepath


def _export_pdf(conv: dict, export_dir: Path, filename: str) -> Path:
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "CampusRule AI - Conversation Export", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, f"Title: {conv.get('title', 'Untitled')}", ln=True)
        pdf.cell(0, 8, f"Date: {conv.get('created_at', '')}", ln=True)
        pdf.ln(5)

        for msg in conv.get("messages", []):
            role = "Student" if msg["role"] == "user" else "CampusRule AI"
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 8, role, ln=True)
            pdf.set_font("Helvetica", "", 10)
            # Handle unicode by encoding safely
            content = msg["content"].encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 6, content)
            pdf.ln(3)

        filepath = export_dir / f"{filename}.pdf"
        pdf.output(str(filepath))
        return filepath
    except ImportError:
        filepath = export_dir / f"{filename}.txt"
        filepath.write_text(_export_text(conv))
        return filepath
